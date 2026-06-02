import re, csv, sys, os, subprocess, tempfile
from docx import Document

def find_word_file():
    docx = [f for f in os.listdir(".") if f.lower().endswith(".docx")]
    doc  = [f for f in os.listdir(".") if f.lower().endswith(".doc") and not f.lower().endswith(".docx")]
    candidates = docx + doc
    if not candidates:
        print("没有 .docx 或 .doc 文件")
        sys.exit(1)
    if len(candidates) == 1:
        return candidates[0]
    print("找到多个文件：")
    for i, f in enumerate(candidates, 1):
        print(f"  [{i}] {f}")
    return candidates[int(input("选一个序号: ")) - 1]

def convert_doc_to_docx(doc_path):
    tmp = tempfile.mkdtemp()
    subprocess.run([
        "soffice", "--headless", "--convert-to", "docx",
        "--outdir", tmp, os.path.abspath(doc_path)
    ], check=True, capture_output=True)
    converted = os.path.join(tmp, os.path.splitext(doc_path)[0] + ".docx")
    if not os.path.exists(converted):
        raise FileNotFoundError("转换失败，未生成 .docx")
    return converted

def extract_text(doc):
    """提取段落 + 表格中的所有文本（去重）"""
    parts = set()
    for p in doc.paragraphs:
        if p.text.strip():
            parts.add(p.text)
    for table in doc.tables:
        for row in table.rows:
            for c in row.cells:
                if c.text.strip():
                    parts.add(c.text)
    return "\n".join(parts)

# --- 主流程 ---
if len(sys.argv) > 1:
    filepath = sys.argv[1]
else:
    filepath = find_word_file()

is_doc = filepath.lower().endswith(".doc") and not filepath.lower().endswith(".docx")
if is_doc:
    print(f"检测到 .doc 文件：{filepath}")
    try:
        filepath = convert_doc_to_docx(filepath)
        print("   已通过 LibreOffice 转换为 .docx")
    except (subprocess.CalledProcessError, FileNotFoundError):
        print("   未找到 LibreOffice，请先安装 LibreOffice 或将 .doc 另存为 .docx")
        sys.exit(1)
else:
    print(f"正在读取: {filepath}")

doc = Document(filepath)
full_text = extract_text(doc)

# --- 提取（结果去重） ---
nct       = sorted(set(re.findall(r"(?:NCT|CTR)\d{6,8}", full_text)))
dates     = sorted(set(re.findall(r"\d{4}-\d{2}-\d{2}", full_text)))
endpoints = sorted(set(
    re.findall(r"(?:主要终点|主要目的|主要评价指标)[：:\s]*([^。\n；]{10,200})", full_text)
))

with open("extracted_info.csv", "w", newline="", encoding="utf-8-sig") as f:
    w = csv.writer(f)
    w.writerow(["类型", "提取内容"])
    for n in nct: w.writerow(["注册号", n])
    for d in dates: w.writerow(["日期", d])
    for e in endpoints: w.writerow(["主要终点/目的", e.strip()])

print(f"提取完成：{len(nct)} 注册号 / {len(dates)} 日期 / {len(endpoints)} 主要终点 -> extracted_info.csv")